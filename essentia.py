import streamlit as st
import streamlit.components.v1 as components
import tempfile
import os
import requests
import base64
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
from datetime import datetime
from dotenv import load_dotenv

# Simplified imports - with proper error handling
try:
    from langchain_groq import ChatGroq
    from langchain_community.embeddings import HuggingFaceEmbeddings
    from langchain.prompts import PromptTemplate
    from langchain.chains import RetrievalQA
    LANGCHAIN_AVAILABLE = True
except ImportError as e:
    st.error(f"LangChain components not available: {e}")
    LANGCHAIN_AVAILABLE = False

# Document processing imports
try:
    import PyPDF2
    import docx
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain.schema import Document
    DOC_PROCESSING_AVAILABLE = True
except ImportError as e:
    st.error(f"Document processing libraries not available: {e}")
    DOC_PROCESSING_AVAILABLE = False

# Qdrant imports
try:
    from qdrant_client import QdrantClient
    from qdrant_client.models import Distance, VectorParams
    from langchain_community.vectorstores import Qdrant
    QDRANT_AVAILABLE = True
except ImportError as e:
    st.error(f"Qdrant not available: {e}")
    QDRANT_AVAILABLE = False

# Load environment variables
load_dotenv(override=True)

# Add this function after your imports and before st.set_page_config()

def inject_pwa_components():
    """Inject PWA manifest, meta tags, and service worker registration"""
    
    pwa_code = """
    <!-- PWA Meta Tags -->
    <link rel="manifest" href="/app/static/manifest.json">
    
    <!-- Theme Color -->
    <meta name="theme-color" content="#ff4b4b">
    <meta name="msapplication-TileColor" content="#ff4b4b">
    
    <!-- iOS Meta Tags -->
    <meta name="apple-mobile-web-app-capable" content="yes">
    <meta name="apple-mobile-web-app-status-bar-style" content="black-translucent">
    <meta name="apple-mobile-web-app-title" content="Essential">
    <link rel="apple-touch-icon" href="/app/static/icons/icon-192x192.png">
    
    <!-- Android/Chrome -->
    <meta name="mobile-web-app-capable" content="yes">
    
    <!-- Windows -->
    <meta name="msapplication-starturl" content="/">
    <meta name="msapplication-TileImage" content="/app/static/icons/icon-144x144.png">
    
    <!-- Additional PWA Meta -->
    <meta name="description" content="Enhanced RAG System with OCR for document analysis and image-based questions">
    
    <!-- Service Worker Registration -->
    <script>
        if ('serviceWorker' in navigator) {
            window.addEventListener('load', function() {
                navigator.serviceWorker.register('/app/static/service-worker.js')
                    .then(function(registration) {
                        console.log('✅ Service Worker registered successfully:', registration.scope);
                        
                        // Check for updates
                        registration.addEventListener('updatefound', () => {
                            const newWorker = registration.installing;
                            console.log('🔄 Service Worker update found');
                            
                            newWorker.addEventListener('statechange', () => {
                                if (newWorker.state === 'installed' && navigator.serviceWorker.controller) {
                                    console.log('✨ New Service Worker available');
                                    // Optionally notify user about update
                                }
                            });
                        });
                    })
                    .catch(function(error) {
                        console.log('❌ Service Worker registration failed:', error);
                    });
            });
            
            // Listen for controller change (new service worker activated)
            navigator.serviceWorker.addEventListener('controllerchange', () => {
                console.log('🔄 Service Worker controller changed');
            });
        } else {
            console.log('⚠️ Service Workers not supported in this browser');
        }
        
        // PWA Install prompt handling
        let deferredPrompt;
        
        window.addEventListener('beforeinstallprompt', (e) => {
            console.log('💾 PWA install prompt available');
            e.preventDefault();
            deferredPrompt = e;
            
            // Show install button or notification
            // You could add a custom install button here
        });
        
        window.addEventListener('appinstalled', () => {
            console.log('✅ PWA installed successfully');
            deferredPrompt = null;
        });
        
        // Detect if app is running as PWA
        if (window.matchMedia('(display-mode: standalone)').matches || 
            window.navigator.standalone === true) {
            console.log('📱 Running as PWA');
            document.body.classList.add('pwa-mode');
        }
    </script>
    
    <!-- PWA Styles -->
    <style>
        /* Hide Streamlit branding when running as PWA */
        .pwa-mode #MainMenu,
        .pwa-mode footer,
        .pwa-mode header {
            visibility: hidden;
        }
        
        /* Optimize for mobile PWA */
        @media (display-mode: standalone) {
            body {
                -webkit-user-select: none;
                -webkit-tap-highlight-color: transparent;
                -webkit-touch-callout: none;
            }
        }
        
        /* Safe area for notched devices */
        @supports (padding: max(0px)) {
            .main {
                padding-left: max(12px, env(safe-area-inset-left));
                padding-right: max(12px, env(safe-area-inset-right));
            }
        }
    </style>
    """
    
    # Use st.components to inject HTML
    components.html(pwa_code, height=0)

# Configure page
st.set_page_config(
    page_title="Essential",
    page_icon="📖",
    layout="centered",
    initial_sidebar_state="expanded"
)

def get_api_key(key_name: str) -> str:
    """Get API key from environment variables or Streamlit secrets"""
    env_value = os.getenv(key_name)
    if env_value and env_value.strip():
        return env_value.strip()
    
    try:
        if hasattr(st, 'secrets') and key_name in st.secrets:
            return str(st.secrets[key_name]).strip()
    except:
        pass
    return ""

# Get API keys
groq_key = get_api_key("GROQ_API_KEY")
qdrant_url = get_api_key("QDRANT_URL")
qdrant_api_key = get_api_key("QDRANT_API_KEY")
ocr_api_key = get_api_key("OCR_SPACE_API_KEY")  # OCR.space API key

# Check dependencies first
if not LANGCHAIN_AVAILABLE or not DOC_PROCESSING_AVAILABLE or not QDRANT_AVAILABLE:
    st.error("Missing required dependencies. Please install:")
    st.code("""
pip install sentence-transformers
pip install langchain-groq
pip install langchain-community
pip install langchain
pip install PyPDF2
pip install python-docx
pip install qdrant-client
pip install requests
pip install python-dotenv
    """)
    st.stop()

# Validate required keys
missing_keys = []
if not groq_key:
    missing_keys.append("GROQ_API_KEY")
if not qdrant_url:
    missing_keys.append("QDRANT_URL")
if not qdrant_api_key:
    missing_keys.append("QDRANT_API_KEY")
if not ocr_api_key:
    missing_keys.append("OCR_SPACE_API_KEY")

if missing_keys:
    st.error(f"Missing API keys: {', '.join(missing_keys)}")
    st.info("Add these to your .env file or Streamlit secrets")
    st.stop()

@dataclass
class ChatMessage:
    role: str
    content: str
    timestamp: datetime
    sources: Optional[List[str]] = None
    query_type: Optional[str] = None  # "document" or "image"

class OCRProcessor:
    """Handles OCR processing using OCR.space API"""
    
    def __init__(self, api_key: str):
        self.api_key = api_key
        self.ocr_url = "https://api.ocr.space/parse/image"
    
    def extract_text_from_image(self, image_path: str) -> str:
        """Extract text from image using OCR.space API"""
        try:
            # Prepare the payload
            payload = {
                'apikey': self.api_key,
                'language': 'eng',
                'isOverlayRequired': False,
                'detectOrientation': True,
                'scale': True,
                'OCREngine': 2,  # Use OCR Engine 2 for better accuracy
            }
            
            # Read and encode the image file
            with open(image_path, 'rb') as image_file:
                files = {'file': image_file}
                
                # Make the API request
                response = requests.post(self.ocr_url, files=files, data=payload, timeout=30)
                
                if response.status_code == 200:
                    result = response.json()
                    
                    # Check if the API call was successful
                    if result.get('IsErroredOnProcessing', True):
                        error_message = result.get('ErrorMessage', ['Unknown error'])
                        if isinstance(error_message, list):
                            error_message = ', '.join(error_message)
                        return f"OCR processing error: {error_message}"
                    
                    # Extract text from the response
                    parsed_results = result.get('ParsedResults', [])
                    if not parsed_results:
                        return "No text could be extracted from the image. Please ensure the image contains clear, readable text."
                    
                    # Combine text from all parsed results
                    extracted_text = ""
                    for parsed_result in parsed_results:
                        text = parsed_result.get('ParsedText', '').strip()
                        if text:
                            extracted_text += text + "\n"
                    
                    extracted_text = extracted_text.strip()
                    
                    if not extracted_text:
                        return "No text could be extracted from the image. Please ensure the image contains clear, readable text."
                    
                    return extracted_text
                    
                else:
                    return f"OCR API request failed with status code: {response.status_code}"
                    
        except requests.exceptions.Timeout:
            return "OCR request timed out. Please try again with a smaller image."
        except requests.exceptions.RequestException as e:
            return f"OCR request failed: {str(e)}"
        except Exception as e:
            return f"Error processing image: {str(e)}"

class DocumentProcessor:
    """Handles document loading and processing"""
    
    def __init__(self, ocr_processor: OCRProcessor):
        self.ocr_processor = ocr_processor
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            st.error(f"Error processing PDF: {e}")
            return ""
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            st.error(f"Error processing DOCX: {e}")
            return ""
    
    @staticmethod
    def extract_text_from_txt(file_path: str) -> str:
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            st.error(f"Error processing TXT: {e}")
            return ""
    
    def extract_text_from_image(self, file_path: str) -> str:
        """Extract text from image using OCR.space API"""
        return self.ocr_processor.extract_text_from_image(file_path)
    
    def process_uploaded_file(self, uploaded_file, file_type: str = "document") -> tuple[str, str]:
        """Process uploaded file - either document or image"""
        with tempfile.NamedTemporaryFile(delete=False, suffix=uploaded_file.name) as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            tmp_path = tmp_file.name
        
        filename = uploaded_file.name
        file_ext = filename.lower().split('.')[-1]
        
        try:
            if file_type == "image":
                # Process as image for OCR
                if file_ext in ['jpg', 'jpeg', 'png']:
                    text = self.extract_text_from_image(tmp_path)
                else:
                    st.error(f"Unsupported image type: {file_ext}")
                    text = ""
            else:
                # Process as document
                if file_ext == 'pdf':
                    text = self.extract_text_from_pdf(tmp_path)
                elif file_ext == 'docx':
                    text = self.extract_text_from_docx(tmp_path)
                elif file_ext == 'txt':
                    text = self.extract_text_from_txt(tmp_path)
                else:
                    st.error(f"Unsupported file type: {file_ext}")
                    text = ""
            
            return filename, text
        finally:
            os.unlink(tmp_path)

class VectorStoreManager:
    """Simplified vector store management"""
    
    def __init__(self):
        self.vectorstore = None
        self.qdrant_client = None
        self.collection_name = "essentia_docs"
        
        # Text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len
        )
        
        # Initialize embeddings with error handling
        try:
            self.embeddings = HuggingFaceEmbeddings(
                model_name="sentence-transformers/all-MiniLM-L6-v2"
            )
        except Exception as e:
            st.error(f"Failed to initialize embeddings: {e}")
            st.error("Please install sentence-transformers: pip install sentence-transformers")
            self.embeddings = None
        
        if self.embeddings:
            self._initialize_qdrant()
            self._load_existing_vectorstore()
    
    def _initialize_qdrant(self):
        """Initialize Qdrant connection"""
        try:
            self.qdrant_client = QdrantClient(url=qdrant_url, api_key=qdrant_api_key)
            collections = self.qdrant_client.get_collections()
            
            # Check if collection exists
            collection_exists = any(c.name == self.collection_name for c in collections.collections)
            
            if not collection_exists:
                self.qdrant_client.create_collection(
                    collection_name=self.collection_name,
                    vectors_config=VectorParams(size=384, distance=Distance.COSINE)
                )
                
        except Exception as e:
            st.error(f"Qdrant connection failed: {e}")
            self.qdrant_client = None
    
    def _load_existing_vectorstore(self):
        """Load existing vectorstore if available"""
        try:
            if self.qdrant_client and self.embeddings:
                self.vectorstore = Qdrant(
                    client=self.qdrant_client,
                    collection_name=self.collection_name,
                    embeddings=self.embeddings
                )
                
                collection_info = self.qdrant_client.get_collection(self.collection_name)
                doc_count = collection_info.points_count
                
                if doc_count > 0:
                    st.session_state.documents_loaded = True
                    st.session_state.document_count = doc_count
                    
        except Exception as e:
            st.warning(f"Could not load existing documents: {e}")
    
    def add_documents(self, documents: List[Document]) -> bool:
        """Add documents to vector store"""
        try:
            if not self.qdrant_client or not self.embeddings:
                st.error("Vector store not properly initialized")
                return False
            
            # Split documents
            chunks = self.text_splitter.split_documents(documents)
            if not chunks:
                st.error("No text chunks created")
                return False
            
            # Create or update vectorstore
            if self.vectorstore is None:
                self.vectorstore = Qdrant.from_documents(
                    chunks,
                    self.embeddings,
                    url=qdrant_url,
                    api_key=qdrant_api_key,
                    collection_name=self.collection_name,
                    force_recreate=False
                )
            else:
                self.vectorstore.add_documents(chunks)
            
            st.success(f"Added {len(chunks)} document chunks")
            return True
            
        except Exception as e:
            st.error(f"Error adding documents: {e}")
            return False
    
    def similarity_search(self, query: str, k: int = 5) -> List[Document]:
        """Search for similar documents"""
        if self.vectorstore:
            try:
                return self.vectorstore.similarity_search(query, k=k)
            except Exception as e:
                st.error(f"Search error: {e}")
        return []
    
    def clear_documents(self) -> bool:
        """Clear all documents"""
        try:
            if self.qdrant_client:
                self.qdrant_client.delete_collection(self.collection_name)
                self.qdrant_client.create_collection(
                    collection_name=self.collection_name,
                    vectors_config=VectorParams(size=384, distance=Distance.COSINE)
                )
                self.vectorstore = None
                st.success("Cleared all documents")
                return True
        except Exception as e:
            st.error(f"Error clearing documents: {e}")
        return False

class SimplifiedRAGSystem:
    """Simplified RAG system without CrewAI - using direct LangChain"""
    
    def __init__(self, vector_store_manager: VectorStoreManager):
        self.vector_store = vector_store_manager
        self.llm = None
        self._initialize_llm()
    
    def _initialize_llm(self):
        """Initialize the language model"""
        try:
            self.llm = ChatGroq(
                groq_api_key=groq_key,
                model="llama-3.1-8b-instant",
                temperature=0.7,
                max_tokens=1500
            )
            
        except Exception as e:
            st.error(f"LLM initialization failed: {e}")
    
    def _create_domain_specific_prompt(self, query: str, mode: str, query_type: str = "document") -> PromptTemplate:
        """Create domain-specific prompts based on query content"""
        
        query_lower = query.lower()
        
        # Detect domain based on keywords
        is_data_vault = any(keyword in query_lower for keyword in 
                           ['data vault', 'hub', 'link', 'satellite', 'business key', 'staging'])
        is_airflow = any(keyword in query_lower for keyword in 
                        ['airflow', 'dag', 'pipeline', 'workflow', 'schedule'])
        is_vaultspeed = any(keyword in query_lower for keyword in 
                           ['vaultspeed', 'automation', 'code generation'])
        
        # Create appropriate system prompt
        if is_data_vault:
            system_role = "You are a Data Vault 2.0 expert with deep knowledge of hubs, links, satellites, and data modeling best practices."
        elif is_airflow:
            system_role = "You are an Apache Airflow specialist with expertise in DAG design, orchestration, and pipeline management."
        elif is_vaultspeed:
            system_role = "You are a VaultSpeed automation expert specializing in code generation and deployment strategies."
        else:
            system_role = "You are a technical expert who can analyze and explain complex documentation clearly."
        
        # Special handling for multiple choice questions from images
        if query_type == "image":
            system_role += " You are answering a multiple choice question that was extracted from an image using OCR."
            
            # Check if this looks like a multiple choice question
            has_choices = any(pattern in query_lower for pattern in [' a)', ' b)', ' c)', ' d)', 'a.', 'b.', 'c.', 'd.', 'a -', 'b -', 'c -', 'd -'])
            
            if has_choices:
                instruction = """This is a multiple choice question. Please:

1. **Identify the main question** clearly
2. **Analyze each choice (A, B, C, D)** individually against the context
3. **Determine which choices are correct** - there may be multiple correct answers
4. **Provide your final answer** in this format:
   - **Correct Answer(s): [Letter(s)]**
   - **Explanation:** Brief explanation of why these are correct and others are wrong

Use the context from the documents to support your analysis. Be thorough but concise."""
            else:
                instruction = """Provide a comprehensive answer using the context from documents. Focus on accuracy and clarity."""
        else:
            if mode == "Overview":
                instruction = """Provide a concise overview with:
• Key points in bullet format
• Brief explanations of important concepts  
• Practical recommendations if applicable
Keep it focused and easy to understand."""
            else:
                instruction = """Provide a detailed analysis with:
• Comprehensive explanation of relevant concepts
• Technical implementation details
• Best practices and recommendations
• Potential challenges and solutions
• Integration considerations
Provide thorough technical depth while maintaining clarity."""
        
        # Add OCR context if this is an image query
        ocr_context = ""
        if query_type == "image":
            ocr_context = "\n\nNote: This question was extracted from an image using OCR, so there might be minor text recognition errors. Please interpret the question in the most logical way."
        
        template = f"""{system_role}

Context from documents:
{{context}}

Question: {{question}}{ocr_context}

Instructions: {instruction}

Answer:"""
        
        return PromptTemplate(
            template=template,
            input_variables=["context", "question"]
        )
    
    def generate_response(self, query: str, mode: str = "Overview", query_type: str = "document") -> tuple[str, List[str]]:
        """Generate response using simple RAG approach"""
        if not self.llm:
            return "LLM not initialized. Please check your Groq API key.", []
        
        if not self.vector_store.vectorstore:
            return "Vector store not initialized. Please check your embeddings setup.", []
        
        try:
            # Get relevant documents
            relevant_docs = self.vector_store.similarity_search(query, k=5)
            
            if not relevant_docs:
                return "No relevant documents found. Please upload documents to the knowledge base first.", []
            
            # Extract sources
            sources = list(set([doc.metadata.get('source', 'Unknown') for doc in relevant_docs]))
            
            # Create domain-specific prompt
            prompt = self._create_domain_specific_prompt(query, mode, query_type)
            
            # Create chain
            chain = RetrievalQA.from_chain_type(
                llm=self.llm,
                chain_type="stuff",
                retriever=self.vector_store.vectorstore.as_retriever(search_kwargs={"k": 5}),
                chain_type_kwargs={"prompt": prompt},
                return_source_documents=True
            )
            
            # Get response
            result = chain({"query": query})
            
            return result["result"], sources
            
        except Exception as e:
            st.error(f"Error generating response: {e}")
            return f"Error: {e}", []

class SimplifiedChatbot:
    """Main chatbot class"""
    
    def __init__(self):
        self.ocr_processor = OCRProcessor(ocr_api_key)
        self.document_processor = DocumentProcessor(self.ocr_processor)
        self.vector_store = VectorStoreManager()
        self.rag_system = SimplifiedRAGSystem(self.vector_store)
    
    def is_ready(self) -> bool:
        """Check if chatbot is ready to use"""
        return (self.vector_store.embeddings is not None and 
                self.rag_system.llm is not None)
    
    def load_documents(self, uploaded_files) -> bool:
        """Load documents into knowledge base"""
        if not uploaded_files or not self.is_ready():
            return False
        
        documents = []
        progress_bar = st.progress(0)
        
        for i, uploaded_file in enumerate(uploaded_files):
            filename, text = self.document_processor.process_uploaded_file(uploaded_file, "document")
            
            if text.strip():
                doc = Document(
                    page_content=text,
                    metadata={"source": filename, "upload_time": datetime.now().isoformat()}
                )
                documents.append(doc)
            
            progress_bar.progress((i + 1) / len(uploaded_files))
        
        if documents:
            success = self.vector_store.add_documents(documents)
            if success:
                st.session_state.documents_loaded = True
                st.session_state.document_count = self._get_document_count()
                return True
        
        return False
    
    def process_image_query(self, uploaded_image) -> tuple[str, str, List[str]]:
        """Process image to extract question and generate answer"""
        if not uploaded_image or not self.is_ready():
            return "", "Chatbot not ready. Please check dependencies and API keys.", []
        
        # Extract text from image
        filename, extracted_text = self.document_processor.process_uploaded_file(uploaded_image, "image")
        
        if not extracted_text or "Error" in extracted_text or "No text could be extracted" in extracted_text or "OCR" in extracted_text:
            return extracted_text, "Could not extract readable text from the image. Please ensure the image contains clear, readable text.", []
        
        # Check if this looks like a multiple choice question and format it better
        formatted_text = self._format_multiple_choice_question(extracted_text)
        
        # Generate response using RAG with special handling for multiple choice
        response, sources = self.rag_system.generate_response(formatted_text, "Overview", "image")
        
        return formatted_text, response, sources
    
    def _format_multiple_choice_question(self, text: str) -> str:
        """Format extracted text to better structure multiple choice questions"""
        import re
        
        # Clean up common OCR issues
        text = text.replace('\n\n', '\n').replace('  ', ' ').strip()
        
        # Try to identify and format multiple choice options
        # Look for patterns like "A)", "A.", "A -", etc.
        patterns = [
            r'([ABCD])\s*\)\s*',  # A) format
            r'([ABCD])\s*\.\s*',  # A. format  
            r'([ABCD])\s*-\s*',   # A - format
            r'([ABCD])\s+',       # A format (with space)
        ]
        
        formatted_text = text
        for pattern in patterns:
            if re.search(pattern, text, re.IGNORECASE):
                # Add line breaks before each choice for better formatting
                formatted_text = re.sub(pattern, r'\n\1) ', formatted_text, flags=re.IGNORECASE)
                break
        
        # Clean up extra whitespace and line breaks
        formatted_text = re.sub(r'\n+', '\n', formatted_text)
        formatted_text = re.sub(r'^\n+', '', formatted_text)
        
        return formatted_text.strip()
    
    def _get_document_count(self) -> int:
        try:
            if self.vector_store.qdrant_client:
                collection_info = self.vector_store.qdrant_client.get_collection(self.vector_store.collection_name)
                return collection_info.points_count
        except:
            pass
        return 0
    
    def clear_documents(self) -> bool:
        success = self.vector_store.clear_documents()
        if success:
            st.session_state.documents_loaded = False
            st.session_state.document_count = 0
        return success
    
    def generate_response(self, query: str, mode: str = "Overview") -> tuple[str, List[str]]:
        if not self.is_ready():
            return "Chatbot not ready. Please check dependencies and API keys.", []
        return self.rag_system.generate_response(query, mode, "document")

def load_image_as_base64(image_path: str) -> str:
    """Load image and convert to base64 - supports both local and git paths"""
    # Define possible paths
    local_path = f"D:\\MOOD\\CODE\\{image_path}"
    git_path = image_path
    
    # Try local path first, then git path
    paths_to_try = [local_path, git_path]
    
    for path in paths_to_try:
        try:
            if os.path.exists(path):
                with open(path, "rb") as img_file:
                    return base64.b64encode(img_file.read()).decode()
        except Exception as e:
            continue
    
    # If no paths work, show warning but don't error
    st.warning(f"Image not found at either: {local_path} or {git_path}")
    return ""

def render_sidebar():
    """Render sidebar controls"""
    with st.sidebar:
        book_img_path = "images/essential.png"
        book_b64 = load_image_as_base64(book_img_path)
        
        if book_b64:
            st.markdown(f"""
                <div style="text-align: center; margin-bottom: 20px;">
                    <img src="data:image/png;base64,{book_b64}" 
                        style="width: 140px; height: auto; margin-bottom: 10px;" />
                    <p style="margin: 0; font-style: italic; color: #666; font-size: 14px;">
                        Enhanced RAG System with OCR.space
                    </p>
                </div>
            """, unsafe_allow_html=True)
        else:
            st.markdown("# 📚 Essential")
            st.markdown("*Powered by Langchain + OCR.space*")
        st.markdown("---")
        
        # Status indicators
        chatbot_ready = st.session_state.get('chatbot') and st.session_state.chatbot.is_ready()
        st.markdown(f"**System Status:** {'✅ Ready' if chatbot_ready else '❌ Not Ready'}")
        st.markdown(f"**Groq LLM:** {'✅ Ready' if groq_key else '❌ Missing'}")
        st.markdown(f"**Qdrant:** {'✅ Connected' if qdrant_url and qdrant_api_key else '❌ Missing'}")
        st.markdown(f"**OCR.space:** {'✅ Ready' if ocr_api_key else '❌ Missing'}")
        st.markdown("---")
        
        # Document management
        st.markdown("### 📄 Knowledge Base")
        
        if st.session_state.get('documents_loaded', False):
            doc_count = st.session_state.get('document_count', 0)
            st.success(f"✅ {doc_count} documents loaded")
            
            if st.button("🗑️ Clear Documents"):
                if st.session_state.chatbot.clear_documents():
                    st.rerun()
        
        uploaded_files = st.file_uploader(
            "📁 Upload Documents (Enhance Knowledge Base)",
            type=['pdf', 'docx', 'txt'],
            accept_multiple_files=True,
            help="Upload your documents to build the knowledge base"
        )
        
        if uploaded_files and chatbot_ready:
            if st.button("📤 Process Documents"):
                with st.spinner("Processing documents..."):
                    if st.session_state.chatbot.load_documents(uploaded_files):
                        st.success("Documents processed!")
                        st.rerun()
        
        st.markdown("---")
        
        # Image question uploader
        st.markdown("### 🖼️ Ask from Image")
        
        uploaded_image = st.file_uploader(
            "🖼️ Upload Image (Ask a Question)",
            type=['jpg', 'jpeg', 'png'],
            help="Upload an image containing a question in text form"
        )
        
        if uploaded_image and chatbot_ready:
            if st.button("🔍 Process Image Question"):
                with st.spinner("Extracting text using OCR.space and generating answer..."):
                    extracted_text, response, sources = st.session_state.chatbot.process_image_query(uploaded_image)
                    
                    # Display extracted text in a text area for user review
                    if extracted_text and not extracted_text.startswith("Error") and not extracted_text.startswith("OCR"):
                        st.markdown("**Extracted Text:**")
                        st.text_area("Text from image:", value=extracted_text, height=100, disabled=True)
                        
                        # Add to chat history
                        if response and not response.startswith("Could not extract"):
                            # Add extracted question as user message
                            user_msg = ChatMessage(
                                role="user", 
                                content=f"[From Image: {uploaded_image.name}] {extracted_text}", 
                                timestamp=datetime.now(),
                                query_type="image"
                            )
                            st.session_state.chat_history.append(user_msg)
                            
                            # Add response as assistant message
                            assistant_msg = ChatMessage(
                                role="assistant", 
                                content=response, 
                                timestamp=datetime.now(),
                                sources=sources,
                                query_type="image"
                            )
                            st.session_state.chat_history.append(assistant_msg)
                            
                            st.success("Image processed and question answered!")
                            st.rerun()
                        else:
                            st.error(response)
                    else:
                        st.error(extracted_text if extracted_text else "Could not extract text from image")
        
        st.markdown("---")
        
        # Settings
        st.markdown("### ⚙️ Settings")
        response_mode = st.selectbox(
            "Response Mode",
            ["Overview", "Deep Dive"],
            help="Choose response detail level"
        )
        
        show_sources = st.toggle("Show Sources", value=True)
        
        return response_mode, show_sources

def main():
    """Main application"""
    
    # Initialize session state
    if 'chat_history' not in st.session_state:
        st.session_state.chat_history = []
    if 'documents_loaded' not in st.session_state:
        st.session_state.documents_loaded = False
    if 'chatbot' not in st.session_state:
        st.session_state.chatbot = SimplifiedChatbot()
    inject_pwa_components()
    # Main header
    quickquery_img_path = "images/essential.png"
    quickquery_b64 = load_image_as_base64(quickquery_img_path)
    
    if quickquery_b64:
        st.markdown(f'''
            <div class="main-header" style="margin-bottom: 20px;">
                <div style="display: flex; align-items: flex-end;">
                    <img src="data:image/png;base64,{quickquery_b64}" 
                        style="width: 240px; height: auto; margin-right: 20px;" />
                    <h1 style="margin: 0; font-size: 15px; font-weight: bold;">
                        Essential turns long documents into concise summaries and can answer questions from images using OCR.space.
                    </h1>
                </div>
            </div>
        ''', unsafe_allow_html=True)
    else:
        st.markdown('''
        <div class="main-header">
            <h1>📖 Essential turns long documents into concise summaries and can answer questions from images using OCR.space.</h1>
        </div>
        ''', unsafe_allow_html=True)
    
    # Check if system is ready
    if not st.session_state.chatbot.is_ready():
        st.error("⚠️ System not ready. Please check dependencies and API keys above.")
        return
    
    # Sidebar
    response_mode, show_sources = render_sidebar()
    
    # Instructions
    st.markdown("### 📋 How to Use")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("""
        **📄 Build Knowledge Base:**
        1. Upload PDF/DOCX/TXT documents in the sidebar
        2. Click "Process Documents" to add them to the knowledge base
        3. Use the chat below to ask questions about uploaded documents
        """)
    
    with col2:
        st.markdown("""
        **🖼️ Ask from Images:**
        1. Upload an image containing a question in the sidebar
        2. Click "Process Image Question" to extract text using OCR.space
        3. Review the extracted text and get an answer from the knowledge base
        """)
    
    # Available topics
    st.markdown("### 📋 Available Topics for Search")
    st.markdown("""
    - Data Vault 2.0 Fundamentals
    - Hubs, Links, and Satellites  
    - Business Keys and Surrogate Keys
    - Staging Layer Best Practices
    - PIT & Bridge Tables
    - Apache Airflow Orchestration
    - VaultSpeed Automation
    """)
    st.markdown("---")
    
    # Chat interface
    for message in st.session_state.chat_history:
        if message.role == "user":
            with st.chat_message("user"):
                if message.query_type == "image":
                    st.markdown("🖼️ **Question from Image:**")
                st.write(message.content)
        else:
            with st.chat_message("assistant"):
                if message.query_type == "image":
                    st.markdown("🤖 **Answer from Knowledge Base:**")
                st.write(message.content)
                if show_sources and message.sources:
                    st.markdown("**Sources:**")
                    for source in message.sources:
                        st.markdown(f"• {source}")
    
    # Text input for direct questions
    if st.session_state.documents_loaded:
        if user_input := st.chat_input("Ask about your documents..."):
            # Add user message
            user_msg = ChatMessage(role="user", content=user_input, timestamp=datetime.now(), query_type="document")
            st.session_state.chat_history.append(user_msg)
            
            # Generate response
            with st.spinner("Analyzing..."):
                response, sources = st.session_state.chatbot.generate_response(user_input, response_mode)
            
            # Add assistant message
            assistant_msg = ChatMessage(
                role="assistant", 
                content=response, 
                timestamp=datetime.now(),
                sources=sources if show_sources else None,
                query_type="document"
            )
            st.session_state.chat_history.append(assistant_msg)
            
            st.rerun()
    else:
        st.info("Upload documents in the sidebar to start asking questions, or upload an image with a question!")
    
    # Clear chat
    if st.session_state.chat_history:
        if st.button("🗑️ Clear Chat"):
            st.session_state.chat_history = []
            st.rerun()

if __name__ == "__main__":
    main()


